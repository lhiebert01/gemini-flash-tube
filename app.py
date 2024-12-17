# Part 1: Imports and Initial Setup

import streamlit as st
import os
from dotenv import load_dotenv
import google.generativeai as genai
from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api.formatters import TextFormatter
from youtube_transcript_api._errors import TranscriptsDisabled, NoTranscriptFound
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
import requests
from bs4 import BeautifulSoup
from urllib.parse import urlparse, parse_qs

# Must be the first Streamlit command
st.set_page_config(
    page_title="Gemini Flash YouTube Video Summary App",
    page_icon="🎯",
    layout="wide"
)

# Session state variables
if "current_summary" not in st.session_state:
    st.session_state.current_summary = None
if "word_doc_binary" not in st.session_state:
    st.session_state.word_doc_binary = None
if "video_processed" not in st.session_state:
    st.session_state.video_processed = False
if "current_transcript" not in st.session_state:
    st.session_state.current_transcript = None
if "current_video_id" not in st.session_state:
    st.session_state.current_video_id = None
if "current_video_title" not in st.session_state:
    st.session_state.current_video_title = None
if "qa_history" not in st.session_state:
    st.session_state.qa_history = []
if "clear_input" not in st.session_state:
    st.session_state.clear_input = False
# Add these with your other session state variables near the top of the file
if "video_count" not in st.session_state:
    st.session_state.video_count = 0
if "query_count" not in st.session_state:
    st.session_state.query_count = 0
if "fast_summary_generated" not in st.session_state:
    st.session_state.fast_summary_generated = False
    

    
# Load environment variables and configure Gemini
load_dotenv(override=True)
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# Constants for prompts
# Update CHUNK_PROMPT for better quote extraction
CHUNK_PROMPT = """Analyze this portion of the video transcript and provide:
1. Key points discussed in this section
2. Notable quotes and spoken content:
   - Capture complete phrases and sentences that:
     * Express important concepts
     * Demonstrate key insights
     * Show speaker's perspective
   - Include speaker identification where possible
   - Include context around the quote
   - Include exact timestamps
3. Technical data, statistics, or numerical information
4. Important concepts and definitions
5. Brief summary of the content

For quotes and technical data:
- Preserve exact wording
- Include timestamps in [HH:MM:SS] format
- Identify speakers when possible
- Include brief context around each quote
- Only include substantive quotes that add value
- Omit filler phrases or incomplete thoughts

Transcript section: """

# Update FINAL_PROMPT for better content organization

FINAL_PROMPT = """Based on the analysis of all sections, provide a comprehensive summary with:
1. Main Topic/Title: Identify the title of the video and the core subject

2. Executive Summary (200 words): Brief overview

3. Key Points (10-20): Most important concepts discussed

4. Detailed Analysis (2000 words): In-depth discussion

5. Notable Quotes and Key Statements:
   - Include only substantive quotes that demonstrate key insights
   - Format: "[HH:MM:SS] Speaker (if known): Quote"
   - Add brief context around each quote
   - Explain significance when relevant
   - Skip if no meaningful quotes are found

6. Technical Data & Statistics: 
   - List all statistical information
   - Skip section if none found

7. Key Terms & Definitions:
   - Technical terminology explained
   - Skip section if none found

8. Concepts & Frameworks:
   - Theoretical frameworks discussed
   - Skip section if none found

9. Timeline & Structure:
   - Content progression
   - Skip if not relevant

10. Practical Applications:
    - Real-world examples
    - Skip if not applicable

Format:
- Include headers only for non-empty sections
- Use clear formatting for quotes: "[HH:MM:SS] Speaker: Quote (Context/Significance)"
- Maintain consistent formatting throughout
- Skip any sections where no relevant content is found

Please synthesize this complete summary: """

FAST_SUMMARY_PROMPT = """Provide a concise executive summary of the video transcript. 
- Limit the summary to 200-500 words. 
- Focus on the main ideas, key insights, and central theme. 
- Skip detailed quotes, technical terms, or extensive sections."""



# Part 2: URL and Transcript Processing Functions

def get_youtube_video_id(url):
    """
    Extract video ID from various YouTube URL formats including mobile versions.
    Handles standard, shortened, mobile, and embedded URLs.
    """
    try:
        # Clean the URL first
        url = url.strip()
        
        # Handle empty or invalid URLs
        if not url:
            return None
            
        # Common YouTube URL patterns
        patterns = [
            # Standard desktop URLs
            r'(?:https?:\/\/)?(?:www\.)?youtube\.com\/watch\?v=([a-zA-Z0-9_-]+)',
            # Mobile URLs
            r'(?:https?:\/\/)?(?:www\.)?m\.youtube\.com\/watch\?v=([a-zA-Z0-9_-]+)',
            # Shortened URLs
            r'(?:https?:\/\/)?(?:www\.)?youtu\.be\/([a-zA-Z0-9_-]+)',
            # Embedded URLs
            r'(?:https?:\/\/)?(?:www\.)?youtube\.com\/embed\/([a-zA-Z0-9_-]+)',
            # Mobile app sharing URLs
            r'(?:https?:\/\/)?(?:www\.)?youtube\.com\/shorts\/([a-zA-Z0-9_-]+)'
        ]
        
        # Try each pattern
        for pattern in patterns:
            match = re.search(pattern, url)
            if match:
                return match.group(1)
                
        # If no pattern matches, try parsing the URL
        parsed_url = urlparse(url)
        
        # Handle standard youtube.com URLs
        if parsed_url.hostname in ['www.youtube.com', 'youtube.com', 'm.youtube.com']:
            query_params = parse_qs(parsed_url.query)
            if 'v' in query_params:
                return query_params['v'][0]
                
        # Handle youtu.be URLs
        if parsed_url.hostname == 'youtu.be':
            return parsed_url.path.lstrip('/')
            
    except Exception as e:
        st.warning(f"Error parsing YouTube URL: {str(e)}")
    
    return None

def get_youtube_title(video_id):
    """Get YouTube video title with enhanced error handling."""
    try:
        url = f"https://www.youtube.com/watch?v={video_id}"
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        response = requests.get(url, headers=headers)
        soup = BeautifulSoup(response.text, 'html.parser')
        title = soup.find('meta', property='og:title')
        return title['content'] if title else None
    except Exception as e:
        st.warning(f"Could not extract video title: {str(e)}")
        return None

def extract_transcript(video_id):
    """
    Enhanced transcript extraction with multiple fallback methods and improved error handling.
    """
    try:
        # First attempt: Default transcript
        try:
            transcript_list = YouTubeTranscriptApi.get_transcript(video_id)
        except Exception as e:
            # Second attempt: Try with auto-generated transcripts
            try:
                transcript_list = YouTubeTranscriptApi.get_transcript(video_id, languages=['en', 'en-US', 'en-GB', 'auto'])
            except Exception as inner_e:
                # Third attempt: Try with manual transcripts
                try:
                    available_transcripts = YouTubeTranscriptApi.list_transcripts(video_id)
                    transcript_list = available_transcripts.find_transcript(['en']).fetch()
                except Exception as final_e:
                    # If all attempts fail, raise a user-friendly error
                    st.error("❌ Unable to access video transcript")
                    st.info("""
                    💡 This could be because:
                    - The video doesn't have subtitles/closed captions enabled
                    - The video is using embedded/burned-in subtitles
                    - The video might be age-restricted or private
                    
                    Try these solutions:
                    1. Choose a similar video that has closed captions enabled
                    2. Check if the video has manual captions available
                    3. Verify the video is publicly accessible
                    """)
                    return None

        # If we got the transcript, format it
        formatted_transcript = ""
        for item in transcript_list:
            timestamp = int(item["start"])
            time_str = f"[{timestamp//3600:02d}:{(timestamp%3600)//60:02d}:{timestamp%60:02d}]"
            formatted_transcript += f"{time_str} {item['text']} "
            
        return formatted_transcript.strip()
        
    except Exception as e:
        st.error(f"❌ Error accessing video: {str(e)}")
        st.info("""
        💡 Please verify:
        - The URL is correct
        - The video is publicly available
        - You're using a supported URL format
        
        Supported formats:
        - youtube.com/watch?v=...
        - youtu.be/...
        - m.youtube.com/watch?v=...
        """)
        return None
    
def process_video_url(youtube_link):
    """
    Process YouTube URL with enhanced error handling and user feedback.
    Returns video_id and transcript if successful, (None, None) if not.
    """
    if not youtube_link:
        st.warning("⚠️ Please enter a YouTube URL")
        return None, None
        
    # Clean the URL
    youtube_link = youtube_link.strip()
    
    # Extract video ID
    video_id = get_youtube_video_id(youtube_link)
    if not video_id:
        st.error("❌ Invalid YouTube URL format")
        st.info("💡 Please make sure to use a valid YouTube URL.\n" +
                "Supported formats:\n" +
                "- youtube.com/watch?v=...\n" +
                "- youtu.be/...\n" +
                "- m.youtube.com/watch?v=...")
        return None, None
        
    # Get video title
    video_title = get_youtube_title(video_id)
    if not video_title:
        st.warning("⚠️ Could not retrieve video title, but proceeding with analysis...")
    
    # Get transcript
    transcript = extract_transcript(video_id)
    if not transcript:
        return None, None
        
    return video_id, transcript

# Part 3: Text Processing and Content Generation Functions

def chunk_text(text, chunk_size=10000):
    """Split text into manageable chunks."""
    words = text.split()
    chunks = []
    current_chunk = []
    current_length = 0
    
    for word in words:
        if current_length + len(word) + 1 <= chunk_size:
            current_chunk.append(word)
            current_length += len(word) + 1
        else:
            chunks.append(' '.join(current_chunk))
            current_chunk = [word]
            current_length = len(word)
    
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    return chunks

def format_response(response_text):
    """Clean up and standardize formatting of Gemini responses."""
    # Replace multiple newlines with double newline
    cleaned = re.sub(r'\n{3,}', '\n\n', response_text)
    
    # Ensure table headers are properly formatted
    cleaned = re.sub(r'\|\s*\n\s*\|', '|\n|---', cleaned)
    
    # Ensure consistent bullet point formatting
    cleaned = re.sub(r'(?m)^\s*[-•]\s*', '- ', cleaned)
    
    return cleaned

def generate_content(text, prompt, retry_count=3):
    """Generate content using Gemini with enhanced error handling and content validation."""
    formatted_prompt = f"""
    {prompt}
    
    Important instructions:
    - Provide detailed, specific content
    - Include actual examples and quotes when available
    - Skip sections that have no meaningful content
    - Maintain proper formatting throughout
    - Do not generate placeholder or filler content
    
    Content to analyze: {text}
    """
    
    for attempt in range(retry_count):
        try:
            model = genai.GenerativeModel("gemini-1.5-flash-latest",
                                        generation_config={
                                            'temperature': 0.7,
                                            'top_p': 0.8,
                                            'top_k': 40
                                        })
            response = model.generate_content(formatted_prompt)
            
            # Format and validate the response
            formatted_response = format_response(response.text)
            
            # Check if response has actual content
            if not re.search(r'[A-Za-z]{50,}', formatted_response):
                if attempt < retry_count - 1:
                    continue  # Try again if response is too short/empty
            
            return formatted_response
        except Exception as e:
            if "API key not valid" in str(e):
                st.error("❌ Invalid API key. Please check your key and try again.")
                st.stop()
            if attempt == retry_count - 1:
                st.error(f"Error generating content: {str(e)}")
                return None
            continue

def analyze_transcript(transcript):
    """Analyze transcript in chunks with progress tracking."""
    chunks = chunk_text(transcript)
    progress_bar = st.progress(0)
    chunk_analyses = []
    
    for i, chunk in enumerate(chunks):
        st.write(f"Analyzing part {i+1} of {len(chunks)}...")
        progress_bar.progress((i + 1) / len(chunks))
        
        chunk_summary = generate_content(chunk, CHUNK_PROMPT)
        if chunk_summary:
            chunk_analyses.append(chunk_summary)
    
    combined_analysis = "\n\n".join(chunk_analyses)
    return generate_content(combined_analysis, FINAL_PROMPT)

def generate_qa_response(question, transcript, summary):
    """Generate Q&A response."""
    prompt = f"""Based on the video transcript and summary, answer this question:

Question: {question}

Summary:
{summary}

Transcript:
{transcript}

Please provide a specific answer based on the video content."""

    return generate_content(prompt, "")

def create_markdown_download(summary, video_title, video_id, qa_history=None):
    """Create markdown format document with Q&A history."""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    markdown_content = f"""# Video Summary: {video_title}

Generated on: {timestamp}
Video Link: https://youtube.com/watch?v={video_id}

![Thumbnail](http://img.youtube.com/vi/{video_id}/0.jpg)

{summary}

"""
    
    if qa_history and len(qa_history) > 0:
        markdown_content += "\n## Questions & Answers\n\n"
        for qa in qa_history:
            markdown_content += f"**Q: {qa['question']}**\n\n"
            markdown_content += f"A: {qa['answer']}\n\n"
    
    markdown_content += "\n---\nGenerated using YouTube Video Summarizer"
    return markdown_content


# Part 4: Streamlit UI Styling and Initial Setup

def setup_streamlit_ui():
    """Configure Streamlit UI styling and components."""
    
    # Enhanced CSS styling
    st.markdown("""
    <style>
    .main-title {
        background: linear-gradient(120deg, #4285F4, #0F9D58);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-size: 3rem;
        font-weight: 800;
        text-align: center;
        padding: 1rem 0;
        margin-bottom: 0.5rem;
        font-family: 'Google Sans', Arial, sans-serif;
    }

    .subtitle {
        color: #5f6368;
        font-size: 1.3rem;
        text-align: center;
        margin-bottom: 1rem;
        font-family: 'Google Sans', Arial, sans-serif;
    }

    .stButton > button {
        background: linear-gradient(120deg, #4285F4, #0F9D58);
        color: white;
        border-radius: 30px;
        padding: 0.6rem 2rem;
        font-weight: bold;
        border: none;
        box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        transition: all 0.3s ease;
    }

    .video-container {
        background: white;
        border-radius: 15px;
        padding: 1.5rem;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
        margin: 2rem auto;
        max-width: 800px;
    }

    .section-header {
        color: #4285F4;
        font-size: 1.8rem;
        font-weight: bold;
        margin: 1.5rem 0;
        padding-left: 0.5rem;
        border-left: 4px solid #4285F4;
    }

    .footer {
        text-align: center;
        padding: 1rem 0;
        color: #666;
        font-size: 0.9rem;
        margin-top: 2rem;
        border-top: 1px solid #eee;
    }

    .linkedin-link {
        color: #0077b5;
        text-decoration: none;
    }

    .api-notice {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
        font-size: 0.8rem;
        color: #666;
    }
    </style>
    """, unsafe_allow_html=True)

    # App Header
    st.markdown("""
    <div style="text-align: center; padding: 1rem;">
        <h1 class="main-title">🎯 Gemini Flash YouTube Video Summary App</h1>
        <p class="subtitle">AI-Powered Video Analysis & Summarization</p>
        <div style="margin-top: 1rem;">
            <a href="https://www.linkedin.com/in/lindsayhiebert/" class="linkedin-link">By Lindsay Hiebert</a>
        </div>
    </div>
    """, unsafe_allow_html=True)


def clean_table_content(text):
    """Clean table content by removing unwanted dashes and formatting."""
    # Remove leading dashes/hyphens
    cleaned = re.sub(r'^[-\s]*', '', text.strip())
    # Remove trailing dashes/hyphens
    cleaned = re.sub(r'[-\s]*$', '', cleaned)
    # Remove standalone dashes
    cleaned = re.sub(r'^---$', '', cleaned)
    # Replace multiple spaces with single space
    cleaned = re.sub(r'\s+', ' ', cleaned)
    return cleaned.strip()

def convert_markdown_to_word(doc, text):
    """Convert markdown formatted text to proper Word formatting."""
    lines = text.split('\n')
    current_table_text = []
    in_table = False

    for line in lines:
        # Handle table lines
        if line.strip().startswith('|'):
            current_table_text.append(line)
            in_table = True
            continue
        elif in_table and line.strip():
            current_table_text.append(line)
            continue
        elif in_table and not line.strip():
            # Process completed table
            if current_table_text:
                # Create table
                rows = []
                for table_line in current_table_text:
                    if not re.match(r'^[\|\s\-:]+$', table_line):  # Skip separator lines
                        cells = [clean_table_content(cell) for cell in table_line.strip('|').split('|')]
                        if any(cell.strip() for cell in cells):  # Only add rows with content
                            rows.append(cells)
                
                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Table Grid'
                    for i, row in enumerate(rows):
                        for j, cell in enumerate(row):
                            table.cell(i, j).text = cell.strip()
                            if i == 0:  # Make header row bold
                                table.cell(i, j).paragraphs[0].runs[0].bold = True
                
                doc.add_paragraph()  # Add spacing after table
            current_table_text = []
            in_table = False
            continue

        # Handle headings
        if line.startswith('#'):
            level = len(re.match(r'^#+', line).group())
            text = line.lstrip('#').strip()
            heading = doc.add_heading(text, level)
            heading.runs[0].bold = True
            continue

        # Handle bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.lstrip('- *').strip()
            p = doc.add_paragraph(text, style='List Bullet')
            continue

        # Handle numbered lists
        if re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            p = doc.add_paragraph(text, style='List Number')
            continue

        # Handle bold text
        if '**' in line:
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:  # Odd parts are bold
                    run.bold = True
            continue

        # Handle regular paragraphs
        if line.strip():
            doc.add_paragraph(line.strip())

    return doc

# Update create_word_document function
def create_word_document(summary, video_title, video_id, qa_history=None):
    """Create Word document with proper formatting."""
    doc = Document()
    
    # Add title
    title = doc.add_heading(f"Video Summary: {video_title or 'Untitled'}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Video Link: https://youtube.com/watch?v={video_id}")
    doc.add_paragraph()
    
    # Convert main content
    doc = convert_markdown_to_word(doc, summary)
    
    # Add Q&A section
    if qa_history and len(qa_history) > 0:
        doc.add_heading("Questions & Answers", 1)
        for qa in qa_history:
            # Question
            q_para = doc.add_paragraph()
            q_run = q_para.add_run("Q: ")
            q_run.bold = True
            q_para.add_run(qa['question'])
            
            # Answer
            a_para = doc.add_paragraph()
            a_run = a_para.add_run("A: ")
            a_run.bold = True
            a_para.add_run(qa['answer'])
            
            doc.add_paragraph()
    
    # Add footer
    doc.add_paragraph()
    footer = doc.add_paragraph("Generated using YouTube Video Summarizer")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

def setup_api_section():
    """Setup API key configuration."""
    # Check if limits are exceeded
    if st.session_state.video_count >= 3 or st.session_state.query_count >= 5:
        st.error("🛑 Session limits reached!")
        st.warning("""
        ### You've reached the free usage limits:
        - Videos analyzed: {}/3
        - Questions asked: {}/5
        """.format(st.session_state.video_count, st.session_state.query_count))
        
        # Create two columns for the choices
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("End Session"):
                st.session_state.clear()  # Clear session state
                st.rerun()
        
        with col2:
            continue_with_key = st.button("Continue with API Key")
            
        if continue_with_key:
            st.info("""
            To continue using the app:
            1. Get your API key from [Google AI Studio](https://aistudio.google.com/apikey)
            2. Enter your key below
            """)
            
            user_api_key = st.text_input("Enter your Google AI API Key:", type="password")
            if user_api_key:
                # Reset counters when new API key is provided
                st.session_state.video_count = 0
                st.session_state.query_count = 0
                return user_api_key
            else:
                st.stop()
        else:
            st.stop()
    
    try:
        # Try to get API key from environment first
        api_key = os.getenv("GOOGLE_API_KEY")
        if not api_key:
            # If not in environment, try Streamlit secrets
            api_key = st.secrets["GOOGLE_API_KEY"]
        return api_key
    except Exception:
        st.error("❌ API key not found")
        st.stop()


def show_quick_guide():
    """Display the Quick Guide section."""
    with st.expander("ℹ️ How to Use"):
        st.markdown("""
        1. 🔑 Enter your Google AI API key
        2. 🔗 Paste any YouTube URL format:
           - Regular: youtube.com/watch?v=...
           - Mobile: m.youtube.com/watch?v=...
           - Short: youtu.be/...
        3. 🚀 Click 'Generate Detailed Notes'
        4. 📝 Get AI-powered summary and insights
        5. ❓ Ask questions about the content
        6. 📥 Download in Markdown or Word format
        """)

def show_footer():
    """Display the app footer."""
    st.markdown("---")  # Horizontal line
    st.markdown("""
    <div style="text-align: center; font-size: 0.8rem; color: #666; margin: 10px 0;">
        End of Analysis
    </div>
    <div style="text-align: center; font-size: 0.7rem; color: #888; margin: 5px 0;">
        Made by Lindsay Hiebert powered by GenAI tools: ❤️ Google Gemini Flash and Streamlit<br>
        <a href="https://www.linkedin.com/in/lindsayhiebert/" style="color: #0077b5; text-decoration: none;">Connect with me on LinkedIn</a>
    </div>
    <div style="text-align: center; font-size: 0.6rem; color: #999; margin-top: 5px;">
        Gemini Flash YouTube Video Summary App
    </div>
    """, unsafe_allow_html=True)
    st.markdown("---")
    
def handle_video_analysis(video_id, transcript):
    """Process video analysis and display results."""

    if st.session_state.video_count >= 3:
        st.error("🛑 Video analysis limit reached (3 videos per session)")
        return

    # Display video thumbnail and title
    st.markdown('<div class="video-container">', unsafe_allow_html=True)
    st.image(f"http://img.youtube.com/vi/{st.session_state.current_video_id}/0.jpg",
             use_container_width=True,
             caption=st.session_state.current_video_title or "Video Thumbnail")
    if st.session_state.current_video_title:
        st.markdown(f'<div style="text-align: center;">{st.session_state.current_video_title}</div>', 
                    unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    # Generate buttons
    col1, col2 = st.columns(2)

    # FAST SUMMARY BUTTON
    with col1:
        if st.button("⚡ Fast Video Summary", key="fast_summary_button"):
            if not st.session_state.fast_summary_generated:
                st.session_state.video_count += 1
                with st.spinner("⏳ Generating fast summary..."):
                    summary = generate_content(transcript, FAST_SUMMARY_PROMPT)
                    if summary:
                        st.session_state.current_summary = summary
                        st.session_state.fast_summary_generated = True
                        st.session_state.video_processed = False
                        st.session_state.qa_history = []  # Reset Q&A history
                        st.rerun()

    # DETAILED SUMMARY BUTTON
    with col2:
        if st.button("🚀 Generate Detailed Notes", key="detailed_summary_button"):
            if not st.session_state.video_processed:
                st.session_state.video_count += 1
                with st.spinner("🔄 Analyzing video content..."):
                    summary = analyze_transcript(transcript)
                    if summary:
                        st.session_state.current_summary = summary
                        st.session_state.video_processed = True
                        st.session_state.fast_summary_generated = False
                        st.session_state.word_doc_binary = None  # Reset word doc binary
                        st.session_state.qa_history = []  # Reset Q&A history
                        st.rerun()


def handle_results_display():
    """Display analysis results and download options."""
    if st.session_state.current_summary:
        # Display success message
        st.success("✨ Summary generated successfully!")
        
        # Display summary header and content
        if st.session_state.video_processed:
            st.markdown('<h2 class="section-header">📋 Detailed Notes</h2>', unsafe_allow_html=True)
        elif st.session_state.fast_summary_generated:
            st.markdown('<h2 class="section-header">⚡ Fast Video Summary</h2>', unsafe_allow_html=True)
        
        st.markdown(st.session_state.current_summary)
        
        # Download options in the main window
        col1, col2 = st.columns(2)
        with col1:
            markdown_content = create_markdown_download(
                st.session_state.current_summary,
                st.session_state.current_video_title,
                st.session_state.current_video_id,
                st.session_state.qa_history
            )
            st.download_button(
                label="📥 Download Markdown",
                data=markdown_content,
                file_name=f"video_summary_{st.session_state.current_video_id}.md",
                mime="text/markdown",
                key="markdown_download_main"
            )
        with col2:
            if st.session_state.word_doc_binary is None:
                doc = create_word_document(
                    st.session_state.current_summary,
                    st.session_state.current_video_title,
                    st.session_state.current_video_id,
                    st.session_state.qa_history
                )
                bio = io.BytesIO()
                doc.save(bio)
                st.session_state.word_doc_binary = bio.getvalue()
            
            st.download_button(
                label="📄 Download Word",
                data=st.session_state.word_doc_binary,
                file_name=f"video_summary_{st.session_state.current_video_id}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="word_download_main"
            )
        
        # Download options in the sidebar
        st.sidebar.markdown("### 📥 Download Summary")
        st.sidebar.download_button(
            label="📥 Download Markdown",
            data=markdown_content,
            file_name=f"video_summary_{st.session_state.current_video_id}.md",
            mime="text/markdown",
            key="markdown_download_sidebar"
        )
        st.sidebar.download_button(
            label="📄 Download Word",
            data=st.session_state.word_doc_binary,
            file_name=f"video_summary_{st.session_state.current_video_id}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="word_download_sidebar"
        )
        
def handle_qa_section():
    """Handle Q&A functionality."""

    if st.session_state.query_count >= 5:
        st.error("🛑 Question limit reached (5 questions per session)")
        return

    # Ensure that a summary (Fast or Detailed) has been generated
    if not st.session_state.current_summary:
        st.warning("⚠️ Generate a summary first to enable Q&A functionality.")
        return

    st.markdown('<h2 class="section-header">❓ Ask Questions About the Video</h2>', 
               unsafe_allow_html=True)

    # Display existing Q&A history
    if st.session_state.qa_history:
        st.markdown("### Previous Questions & Answers")
        for qa in st.session_state.qa_history:
            st.markdown(f"**Q: {qa['question']}**")
            st.markdown(f"A: {qa['answer']}\n")

    # Input box for the user's question
    user_question = st.text_input(
        "Enter your question:",
        key="qa_input"
    )

    # Process user question only if input is not empty and it's a new input
    if user_question and st.session_state.get("last_qa_input", "") != user_question:
        # Save the input to prevent repeated processing
        st.session_state.last_qa_input = user_question

        # Increment query count
        st.session_state.query_count += 1

        # Use the current summary (Fast or Detailed)
        summary = st.session_state.current_summary
        transcript = st.session_state.current_transcript

        with st.spinner("🤔 Analyzing your question..."):
            answer = generate_qa_response(user_question, transcript, summary)
            if answer:
                # Add to Q&A history
                st.session_state.qa_history.append({
                    "question": user_question,
                    "answer": answer
                })

                # Update the Word document binary
                doc = create_word_document(
                    st.session_state.current_summary,
                    st.session_state.current_video_title,
                    st.session_state.current_video_id,
                    st.session_state.qa_history
                )
                bio = io.BytesIO()
                doc.save(bio)
                st.session_state.word_doc_binary = bio.getvalue()

                # Display the latest answer
                st.markdown("### 💡 Latest Answer")
                st.markdown(answer)

                
def show_usage_stats():
    """Display usage statistics."""
    st.sidebar.markdown("### Usage Statistics")
    st.sidebar.text(f"Videos Analyzed: {st.session_state.video_count}/3")  # Changed from 10 to 3
    st.sidebar.text(f"Questions Asked: {st.session_state.query_count}/5")   # Changed from 25 to 5
    if st.session_state.video_count >= 2 or st.session_state.query_count >= 4:  # Changed thresholds
        st.sidebar.warning("⚠️ Approaching usage limit!")
        

def main():
    """Main application function."""
    # Setup UI
    setup_streamlit_ui()
    
    # Add this line to show usage stats
    show_usage_stats()
    
    # Get and validate API key
    api_key = setup_api_section()
    genai.configure(api_key=api_key)
    
    # Show user guide
    show_quick_guide()
    
    # YouTube URL input
    youtube_link = st.text_input("🎥 Enter YouTube Video Link:", 
                                placeholder="https://www.youtube.com/watch?v=... or youtu.be/...")
    
    if youtube_link:
        # Process the URL and get video information
        video_id, transcript = process_video_url(youtube_link)
        
        if video_id and transcript:
            # Only process if it's a new video
            if video_id != st.session_state.current_video_id:
                video_title = get_youtube_title(video_id)
                
                # Store current video information
                st.session_state.current_video_id = video_id
                st.session_state.current_video_title = video_title
                st.session_state.current_transcript = transcript
                # Reset previous results
                st.session_state.current_summary = None
                st.session_state.word_doc_binary = None
                st.session_state.video_processed = False
                st.session_state.qa_history = []
            
            # Handle video analysis
            handle_video_analysis(video_id, transcript)
            
            # Display results if available
            handle_results_display()
            
            # Show Q&A section if either a Fast Summary or Detailed Notes are generated
            if st.session_state.video_processed or st.session_state.fast_summary_generated:
                handle_qa_section()
    
    # Show footer
    show_footer()

if __name__ == "__main__":
    main()